from docx import Document
import os
import json

#Good to check if the working directory is the same as the default
os.getcwd()
#Change if necessary
os.chdir("file/path")

#Check the files in the directory
files = os.listdir(".")
print(files)

#For single file checks
'''Check filetype for a document'''
document = Document("docName.docx")
type(document)
document.paragraphs
#Output will be along the lines of [<docx.text.paragraph.Paragraph object at 0x0000028A60C239D0>

#Create loop to create human readable paragraphs from the document that is used in the function after
index = 0 
for para in document.paragraphs:
  index += 1
  if (len(para.text) > 0:
      print("\n paragraph", index, "is")
      print(para.text)
#Output will be - paragraph 1 is ""
      
def doc_to_dict(fileName):
'''Create function to put text within a dictionary where the paragraph is a key and the text is a value'''
  docx_content = {} 
  document = Document(fileName)
  index = 0
  for para in document.paragraphs:
    index += 1
    if (len(para.text) > 0):
      docx_content[index] = para.text
  return docx_content

#Test the function
doc_to_dict("fileName.docx")

#Optional for multiple files - check all documents of a certain filetype and put them into a dictionary 
docx_content = {}
def files_dir():
  for filename in os.listdir():
    if filename.endswith(".docx"):
      #Replace endswith("") with your file type
      path_docx = os.path.join(filename)
      docx_content[path_docx] = doc_to_dict(path_docx)
      print(os.path.join(filename))
      
#After this point, word cloud can be made by using the word cloud and matplotlib packages, but did not work for me, so I exported to R
#There should be a better way of exporting a dictionary to a text file; let me know 
json.dump(docx_content, file)
file.close()
file = open("jobs.txt", "r")
output = file.read()
print(output)
#the file, in this case is "jobs.txt" should be in your current directory       
      
